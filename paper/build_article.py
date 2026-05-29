# -*- coding: utf-8 -*-
"""Build the new (measured) CitiVerse SDN+DQN article as a .docx.
Keeps the aspirant's file untouched; writes a NEW file in the same folder.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

FIGDIR = r"C:\Users\Максим\Downloads\citiverse_results"
OUT = r"C:\Users\Максим\OneDrive\Рабочий стол\бонч\магистратура\статьи\статья dqn_sdn\dqn simenv+mininet success\paper\Dynamic_clustering_SDN_DQN_CitiVerse_measured.docx"

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Times New Roman'
st.font.size = Pt(12)
doc.styles['Normal'].paragraph_format.space_after = Pt(6)

def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.name = 'Times New Roman'
    return h

def P(text, italic=False, align=None, size=12, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic; r.bold = bold; r.font.size = Pt(size)
    if align == 'c': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'j': p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def fig(fname, caption, width=6.2):
    path = os.path.join(FIGDIR, fname)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = doc.add_paragraph()
    rr = c.add_run(caption); rr.italic = True; rr.font.size = Pt(10)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]; cell.text = ''
        run = cell.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(10)
    return t

# ---------------------------------------------------------------- TITLE
ti = doc.add_paragraph(); ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ti.add_run("Dynamic Clustering of Distributed SDN Controllers for CitiVerse "
               "Networks: A Measured Deep Reinforcement Learning Approach")
r.bold = True; r.font.size = Pt(15)

au = doc.add_paragraph(); au.alignment = WD_ALIGN_PARAGRAPH.CENTER
au.add_run("[Author 1], [Author 2], [Supervisor]\n").font.size = Pt(11)
af = doc.add_paragraph(); af.alignment = WD_ALIGN_PARAGRAPH.CENTER
af.add_run("[The Bonch-Bruevich Saint Petersburg State University of Telecommunications], "
           "Saint Petersburg, Russia\n[e-mail]").italic = True
af.runs[0].font.size = Pt(10)

# ---------------------------------------------------------------- ABSTRACT
H("Abstract", 2)
P("Distributed software-defined networking (SDN) has become a practical way to "
  "manage the dense edge infrastructure of smart-city, or CitiVerse, networks, "
  "where large numbers of cameras, environmental sensors and roadside units "
  "generate control traffic that shifts with the daily rhythm of the city. A "
  "widespread deployment choice maps every switch to the controller of its "
  "administrative zone. This keeps propagation delay low and is easy to operate, "
  "but it cannot follow the diurnal migration of load: during the morning peak the "
  "residential controllers saturate while the commercial and industrial controllers "
  "sit almost idle, and the time needed to install a new flow grows by an order of "
  "magnitude. We cast switch-to-controller assignment as a sequential decision "
  "problem and solve it online with a Dueling Double Deep Q-Network trained with "
  "prioritized experience replay, n-step returns and a staged curriculum. Because "
  "end-to-end training requires on the order of 1.8 million environment steps, the "
  "agent is trained inside a fast analytical queuing model and then transferred to a "
  "live emulation, following a sim-to-real protocol. The central contribution of this "
  "work is that the reported numbers are measured rather than assumed: every latency "
  "value is obtained by timing real OpenFlow flow setups on a Mininet network of "
  "twenty Open vSwitch nodes driven by five independent Ryu controllers under "
  "packet-level load. Averaged over three random seeds, the learned policy lowers the "
  "mean flow-setup latency of the skewed morning profile from 419 ms under "
  "zone-static assignment to 38 ms, while remaining competitive with the strongest "
  "static baseline during the more balanced daytime, evening and night profiles. The "
  "improvement is explained directly by the measurements: the agent spreads the "
  "control load of the busy residential zones across all five controllers instead of "
  "overloading two of them.")

kw = doc.add_paragraph()
kw.add_run("Keywords: ").bold = True
kw.add_run("software-defined networking, distributed SDN controllers, controller "
           "placement, dynamic clustering, deep reinforcement learning, Deep Q-Network, "
           "flow-setup latency, load balancing, smart-city networks, Mininet, OpenFlow.")
for r_ in kw.runs: r_.font.size = Pt(11)

# ---------------------------------------------------------------- 1 INTRO
H("1. Introduction", 1)
P("Smart-city and CitiVerse deployments connect thousands of edge and fog devices, "
  "from traffic signals and surveillance cameras to public access points and "
  "environmental sensors, and are increasingly framed within the 5G/6G smart-city "
  "vision [1]. The traffic these devices produce is dense, bursty and, above all, "
  "non-stationary: its spatial distribution changes over the course of a day as "
  "people move between residential, commercial and industrial districts [2]. The way "
  "such smart-city IoT traffic interacts with an SDN core has been characterised "
  "directly for city segments [3]. Software-defined networking suits these networks "
  "for a concrete reason: by separating the control plane from the data plane and "
  "exposing a programmable, network-wide view, it lets the switch-to-controller "
  "mapping be recomputed and reinstalled at run time instead of being fixed at design "
  "time [4]. Security and management studies of SDN for smart cities point the same "
  "way [5], while the move toward 6G, with its sub-millisecond latency targets and "
  "AI-native control, sharpens the need for this run-time adaptivity [6]. At city "
  "scale a single controller is neither reliable nor fast enough, so the control "
  "plane is distributed across several controllers, each responsible for a subset of "
  "switches and often paired with edge computing [7].")
P("How switches are partitioned among controllers is the controller placement and "
  "clustering problem. The classical formulations are static: switches are grouped "
  "by geographic proximity or by k-means over a delay matrix, and the partition is "
  "fixed at deployment time [8]. A substantial body of work refines this static "
  "partition with learning-automaton [9] and metaheuristic [10] schemes, but the "
  "partition still does not change once the network is running. "
  "Static partitions are attractive operationally and "
  "minimise propagation delay, but they are blind to load. When demand concentrates "
  "in a few zones, the controllers of those zones receive a disproportionate share of "
  "PACKET_IN events, their event queues build up, and the latency experienced when a "
  "new flow is set up rises steeply, even though other controllers remain idle. This "
  "is exactly the regime in which a CitiVerse network spends its mornings and evenings.")
P("Reinforcement learning has been proposed to make placement adaptive [11], and "
  "deep reinforcement learning in particular can handle the large discrete decision "
  "space that arises when any switch may be reassigned to any controller [12]. Most of "
  "the existing literature, however, evaluates such policies on analytical models or "
  "on synthetic parameter sets, and reports performance trends rather than measured "
  "quantities. The practical question — does a learned policy actually reduce the "
  "control-plane latency that an operator would observe on real switches and "
  "controllers — is rarely answered directly.")
P("This paper addresses that question. Our contributions are the following. "
  "(i) We formulate dynamic switch-to-controller assignment for a realistic "
  "twenty-switch, five-controller CitiVerse topology with four diurnal load profiles, "
  "and we adopt flow-setup latency, the operator-visible cost of control-plane "
  "responsiveness, as the primary objective. (ii) We design and train a Dueling "
  "Double Deep Q-Network with prioritized experience replay, n-step returns, action "
  "masking and curriculum learning, and we train it in a fast analytical environment "
  "so that the roughly 1.8 million steps required remain feasible. (iii) We validate "
  "the trained policy on a live Mininet and Ryu emulation, measuring flow-setup "
  "latency from real OpenFlow exchanges under packet-level load, and we compare "
  "against three deployable static baselines over three random seeds. The measured "
  "results show an order-of-magnitude reduction in latency under the most skewed "
  "profile and competitive behaviour elsewhere, and the per-controller measurements "
  "make the mechanism behind the improvement explicit.")

# ---------------------------------------------------------------- 2 SYSTEM MODEL
H("2. System model and problem formulation", 1)
P("We model the CitiVerse access network as a set of OpenFlow switches "
  "S = {s₁, …, sₙ} and a set of distributed controllers "
  "C = {c₁, …, cₖ}. Switches are organised into administrative zones that "
  "reflect the urban layout — residential, commercial and industrial districts — and "
  "each zone is served by default by one controller located in it. Switches inside a "
  "zone are connected in a local chain, and zones are joined by a small number of "
  "higher-delay backbone links. Every switch is attached to its controller over an "
  "OpenFlow control channel with a propagation delay that depends on whether the "
  "controller sits in the same zone or in a remote one.")
P("The assignment is described by a binary matrix A, where aᵢⱼ = 1 if switch "
  "sᵢ is controlled by cⱼ and 0 otherwise. Each switch is assigned to exactly "
  "one controller, so Σⱼ aᵢⱼ = 1 for every i, and a controller is "
  "active when it owns at least one switch. The load offered to a controller is the "
  "sum of the demands of the switches assigned to it, "
  "L(cⱼ) = Σᵢ aᵢⱼ · load(sᵢ), where load(sᵢ) depends "
  "on the zone of sᵢ and on the current traffic profile. Each controller together "
  "with the switches assigned to it forms a cluster, so A fully determines the "
  "clustering; we speak of the assignment throughout and reserve clustering for the "
  "grouping it induces.")
P("Unlike the static formulations, our objective targets control-plane "
  "responsiveness directly. The quantity an operator cares about is the flow-setup "
  "latency: the time from a table-miss at a switch to the installation of the "
  "corresponding flow rule by the controller. This latency has two components. The "
  "first is the propagation delay d(sᵢ, cⱼ) between the switch and its "
  "controller, which a zone-aligned assignment minimises. The second is a queuing "
  "component that grows with controller load, because a controller that receives many "
  "PACKET_IN events processes each one after a queue. We summarise both components in "
  "an inter-controller delay (ICD) term that the agent minimises during training,")
P("ICD(A) = (1/n) Σᵢ [ d(sᵢ, c(sᵢ)) + w(L(c(sᵢ))) ],",
  align='c', italic=True)
P("where c(sᵢ) is the controller assigned to sᵢ and w(·) is a queuing-delay "
  "function that is small at low load and rises sharply as the load approaches the "
  "controller service capacity. Minimising ICD therefore rewards two things at once: "
  "keeping switches close to their controller (low propagation) and avoiding "
  "overloaded controllers (low queuing). The assignment is subject to the usual "
  "single-assignment and controller-activation constraints; we do not impose hard "
  "cluster-size bounds, leaving the agent free to trade propagation against queuing.")
P("It is important to separate this training objective from the way the method is "
  "evaluated. ICD is an analytical surrogate used only to train the agent efficiently "
  "(Section 3). The performance reported in Section 5 is the flow-setup latency "
  "measured on a live emulation, not the value of ICD. We make this distinction "
  "explicit precisely because conflating an analytical objective with measured "
  "performance is a common weakness in the literature.")

# ---------------------------------------------------------------- 3 DQN
H("3. Dynamic clustering with a Deep Q-Network", 1)
P("We treat reassignment as a Markov decision process and learn a policy with a deep "
  "Q-network. At each step the agent observes the network state, selects one "
  "reassignment, and receives a reward derived from the resulting ICD.", )
P("State.", bold=True)
P("The state is a 94-dimensional vector that summarises the current configuration "
  "and load of the network: the normalised controller index of every switch; the "
  "load of the controller each switch is currently attached to; the zone of every "
  "switch and a flag indicating whether the switch is served by the controller of its "
  "own zone; aggregate indicators (normalised ICD, throughput, maximum controller "
  "load and the standard deviation of controller load); the per-controller PACKET_IN "
  "rate; and a short encoding of the active traffic profile. This representation gives "
  "the agent both the current mapping and the load imbalance it has to correct.")
P("Action.", bold=True)
P("An action reassigns a single switch to a single controller, so the action space "
  "has n × k = 20 × 5 = 100 discrete actions. Actions that would violate a constraint "
  "are removed with an action mask before the policy selects, which keeps exploration "
  "inside the feasible region and speeds up learning. We deliberately keep the action "
  "space discrete; operations such as adjusting controller capacity are outside the "
  "scope of a value-based agent and are not modelled as continuous actions.")
P("Reward.", bold=True)
P("The reward is r = − ICD / 20 − 0.02 · m, where m is one if the action "
  "actually migrated a switch and zero otherwise. The first term drives the agent "
  "towards low-latency, well-balanced configurations; the small migration penalty "
  "discourages needless churn that would disrupt established flows. Because ICD "
  "already contains both propagation and queuing, a single scalar reward is enough to "
  "express the full objective.")
P("Network and learning algorithm.", bold=True)
P("The value function is approximated by a Dueling network [13] with a shared "
  "256-unit hidden representation and LayerNorm, which separates the state value from "
  "the action advantage and stabilises learning when many actions have similar value. "
  "We use Double Q-learning [14] to reduce the overestimation bias of plain DQN, a "
  "prioritized experience replay buffer [15] (capacity 10⁵, α = 0.6, "
  "importance-sampling β annealed from 0.4 to 1.0) so that informative transitions "
  "are replayed more "
  "often, and 3-step returns to propagate reward faster. The network is trained with "
  "Adam at a learning rate of 3·10⁻⁴ under a cosine schedule, a discount "
  "factor of 0.99, mini-batches of 256, and a target network refreshed every 100 "
  "steps. Exploration follows an ε-greedy schedule decaying from 1.0 to 0.05 over "
  "the first 80 % of training. We use a value-based agent rather than a "
  "policy-gradient method such as PPO or SAC for reasons that fit the problem: the "
  "action space is discrete (a switch moves to one of five controllers), an "
  "off-policy replay buffer reuses the comparatively expensive transitions far more "
  "efficiently, and the greedy policy used at deployment is deterministic, which is "
  "preferable for a control function.")
P("Curriculum and sim-to-real training.", bold=True)
P("Training the agent end-to-end on the live emulation is infeasible: three seeds of "
  "3000 episodes at 200 steps each amount to about 1.8 million environment steps, far "
  "beyond what a Mininet-in-the-loop setup can deliver in reasonable time. We "
  "therefore train inside the fast analytical environment of Section 2, in which ICD "
  "is computed from the propagation matrix and a convex, M/M/1-inspired queuing "
  "penalty w(L) = 0.4 · L² / (μ − L), capped at 50 ms, with controller service "
  "rate μ = 14 load units. We deliberately use a quadratic numerator rather than "
  "the linear M/M/1 waiting-time expression: the training signal has to penalise "
  "near-saturation operation more sharply than a linear term would, and the "
  "coefficient was set so that the penalty stays in a realistic millisecond range. "
  "The term is a training surrogate that shapes the policy, not a claim about the "
  "controllers' exact queueing behaviour — that behaviour is measured directly in "
  "Section 5. Learning proceeds along a three-stage curriculum: "
  "the first 300 episodes use a reduced 5-switch, 2-controller topology under the "
  "morning profile; the next 500 episodes grow it to 10 switches and 3 controllers "
  "under the morning and business profiles; and the remaining episodes use the full "
  "20-switch, 5-controller topology under all four profiles. The curriculum lets the "
  "agent master a simple version of the task before facing the full problem, where "
  "exploration has already decayed and the policy exploits what it has learned. The "
  "resulting policy is then frozen and evaluated on the live emulation, a standard "
  "sim-to-real arrangement that keeps the analytical model strictly in the role of a "
  "training environment.")

# ---------------------------------------------------------------- 4 TESTBED
H("4. Experimental testbed and methodology", 1)
P("Topology.", bold=True)
P("The evaluation network, shown in Fig. 1, comprises 20 Open vSwitch instances "
  "running OpenFlow 1.3 [16] in an emulation built with Mininet [17], grouped into "
  "five zones and served by five Ryu controllers "
  "listening on ports 6653–6657. Two residential zones contain six and five switches, "
  "two commercial zones contain four and three, and an industrial zone contains two. "
  "Intra-zone links have a propagation delay of 3 ms; the four inter-zone backbone "
  "links carry delays of 8, 12, 15 and 20 ms. Each switch is connected to two hosts, "
  "giving 40 hosts in total: one host per switch generates load and the other serves "
  "as a measurement probe.")
fig("fig_topology_v2.png", "Fig. 1. CitiVerse testbed: 20 OpenFlow switches in five "
    "zones, five distributed Ryu controllers, intra-zone chains and the inter-zone "
    "backbone (data-plane delays annotated). Dashed lines are OpenFlow control "
    "channels under the zone-static (ZoneOptimal) mapping.")
P("Controllers run a reactive layer-2 learning-switch application. Switches operate "
  "in secure fail mode, so a switch without a matching rule cannot forward on its own: "
  "an unknown flow necessarily triggers a PACKET_IN to the controller, which computes "
  "the output port and installs a FLOW_MOD. The time to set up the first packet of a "
  "flow therefore includes the controller's processing time and grows when the "
  "controller is busy. Because Ryu is single-threaded, an overloaded controller builds "
  "a real event-queue backlog, which is precisely the effect we want to measure.")
P("Load profiles.", bold=True)
P("Four diurnal profiles set the per-zone demand through a multiplier applied to a "
  "base request rate (Table 1). Their diurnal structure mirrors the way smart-city "
  "IoT traffic concentrates and shifts between districts over the day. "
  "The morning profile concentrates demand in the "
  "residential zones, the business profile in the commercial zones, the evening "
  "profile spreads it across residential and commercial zones, and the night profile "
  "shifts it to the industrial zone. The morning profile is the most skewed and is the "
  "case in which dynamic clustering has the most to gain.")
table(["Profile", "res1", "res2", "com1", "com2", "ind"],
      [["Morning", 2.5, 2.0, 0.5, 0.5, 0.3],
       ["Business", 0.7, 0.7, 2.5, 2.0, 1.0],
       ["Evening", 2.0, 2.0, 1.5, 1.5, 0.5],
       ["Night", 0.3, 0.3, 0.3, 0.3, 2.5]])
P("Table 1. Per-zone load multipliers for the four diurnal traffic profiles.",
  italic=True, align='c', size=10)
P("Measurement procedure.", bold=True)
P("For each combination of assignment policy and traffic profile we apply the "
  "mapping with ovs-vsctl, load the network, and measure flow-setup latency. Load is "
  "generated by a lightweight raw-socket flooder running on each switch's load host: "
  "it emits frames with a varying source MAC address toward a peer on the same switch, "
  "so every frame is a fresh table-miss and produces a sustained, switch-local stream "
  "of PACKET_IN events to that switch's controller. The base rate is 130 packets per "
  "second, scaled per switch by the zone multiplier of the active profile. To measure "
  "the latency of a switch we clear its flow table, reinstall only the table-miss "
  "entry so that the next flow is genuinely new, and time a ping between the switch's "
  "two hosts; the round-trip to the first reply is the flow-setup latency. Each "
  "measurement is repeated and we report the mean, together with the per-controller "
  "mean and the realised controller loads. Measurement is run after training, on the "
  "freed CPU cores, to avoid contention between the five single-threaded controllers "
  "and the training process.")
P("Baselines.", bold=True)
P("We compare against three deployable static policies. ZoneOptimal assigns each "
  "switch to the controller of its own zone, which is optimal for propagation delay "
  "and is the standard operational default. LoadBalanced distributes switches evenly "
  "across controllers regardless of geography. KMeans clusters switches by their "
  "controller-delay vectors. All three are static and unaware of the active profile; "
  "the learned policy is the only one that adapts online, which is the comparison we "
  "want to make.")
P("Protocol.", bold=True)
P("To report results with a measure of variability rather than a single run, we train "
  "three independent agents with seeds 42, 123 and 456 and report the mean and "
  "standard deviation of the measured latency across them.")

# ---------------------------------------------------------------- 5 RESULTS
H("5. Results and discussion", 1)
P("Convergence.", bold=True)
P("Figure 2 shows the training reward and the analytical ICD, each as a rolling mean "
  "over the three seeds with a one-standard-deviation band, and with the curriculum "
  "stage boundaries marked. The reward drops at each stage transition, as expected "
  "when the task grows harder, and then climbs steadily through the final stage as the "
  "agent learns the full 20-switch problem; the training ICD falls correspondingly. "
  "The curves confirm that the policy converges within the training budget rather than "
  "stalling, and that the three seeds behave consistently.")
fig("fig_training_curves.png", "Fig. 2. Training reward (top) and analytical training "
    "ICD (bottom), rolling mean of three seeds with ± standard-deviation band; "
    "dashed lines mark curriculum stage transitions. The analytical ICD is the "
    "training signal, not a reported result.")
P("Measured flow-setup latency.", bold=True)
P("Table 2 and Fig. 3 report the measured mean flow-setup latency for every policy "
  "and profile. The morning profile is decisive: zone-static assignment overloads the "
  "two residential controllers and the mean latency reaches 419 ms, while the learned "
  "policy keeps it at 38 ms, an order-of-magnitude reduction, and also clearly beats "
  "the load-aware KMeans and LoadBalanced baselines. In the more balanced evening and "
  "night profiles the learned policy is the best or tied-best, although the margins "
  "are small because at light load every policy keeps the controllers below "
  "saturation. The one profile where the agent does not win is business, where the "
  "demand is moderate and falls neatly onto the commercial zones, so the "
  "propagation-minimising ZoneOptimal mapping is already close to ideal; even there "
  "the learned policy comfortably outperforms KMeans.")
table(["Profile", "ZoneOptimal", "LoadBalanced", "KMeans", "DQN (mean ± std)"],
      [["Morning", "418.8", "206.7", "136.9", "38.0 ± 19.0"],
       ["Business", "14.5", "16.2", "89.5", "23.5 ± 4.8"],
       ["Evening", "28.5", "32.7", "26.8", "26.7 ± 3.3"],
       ["Night", "11.7", "13.6", "12.2", "11.5 ± 0.8"]])
P("Table 2. Measured mean flow-setup latency (ms) on the live Mininet/Ryu testbed; "
  "DQN over three seeds. Lower is better.", italic=True, align='c', size=10)
fig("measured_latency_by_profile.png", "Fig. 3. Measured flow-setup latency by "
    "profile and policy (error bar on DQN spans the three seeds). Lower is better.")
P("Why the policy wins.", bold=True)
P("The measurements explain the morning result without appeal to the model. Under "
  "zone-static assignment the two residential controllers receive aggregate loads of "
  "about 15 and 10 units while the other three controllers carry less than 2 units "
  "each; the busy controllers saturate, and the per-controller flow-setup latency "
  "reaches 845 and 1138 ms, whereas the idle controllers stay below 100 ms (Fig. 4). "
  "The learned policy moves a number of residential switches onto the underused "
  "controllers, so that no controller is driven into saturation and the per-controller "
  "latency settles uniformly around 20–30 ms. The throughput-balancing effect is also "
  "visible in the realised controller loads (Fig. 5). In other words, the agent has "
  "discovered control-plane load balancing on its own, and the benefit appears as a "
  "real, measured drop in latency rather than as an assumed trend.")
fig("fig_per_controller_latency.png", "Fig. 4. Per-controller flow-setup latency under "
    "the morning profile. ZoneOptimal saturates C0 and C1; the learned policy keeps "
    "all five controllers near 20–30 ms.")
fig("measured_controller_loads.png", "Fig. 5. Realised controller loads under the "
    "morning profile: the learned policy spreads load that ZoneOptimal concentrates "
    "on two controllers.")
P("A closer look at the business profile.", bold=True)
P("It is worth saying why the agent trails ZoneOptimal here rather than matching it. "
  "Under the business profile the busiest controller carries about ten load units, "
  "comfortably below the service rate, so no controller saturates and the queuing "
  "term is small for every policy. The agent nevertheless applies the load-spreading "
  "behaviour it learned where saturation does occur, moving some switches off their "
  "home controller; this adds propagation delay without buying any queuing reduction, "
  "which is exactly why it falls behind the propagation-optimal static map. The "
  "behaviour is consistent rather than mistaken — the policy is tuned for the "
  "saturated regime that dominates the morning and evening peaks — but it shows that "
  "a single ICD reward does not cleanly separate “balance the load when a controller "
  "is near saturation” from “stay local when there is slack.” Heavier exposure of the "
  "moderate-load profiles during the curriculum, or a reward that explicitly credits "
  "staying close when no controller is loaded, should close the gap.")
P("Limitations.", bold=True)
P("Two points deserve an honest statement. First, the learned policy does not "
  "dominate in every regime: in the business profile a propagation-optimal static map "
  "is already near-ideal and is not improved upon. The contribution is therefore best "
  "described as a large gain under load skew together with competitive behaviour "
  "elsewhere, rather than a uniform improvement. Second, the analytical training "
  "environment and the live measurement do not agree perfectly, especially at light "
  "load; this is expected of any sim-to-real pipeline and is the reason we report only "
  "measured quantities as results and keep the analytical model in its training role. "
  "The spread across seeds under the morning profile is also non-trivial — one seed "
  "settled at 65 ms against 24 ms for the other two — but all three improve on "
  "zone-static assignment by a wide margin. We also note that three seeds is at the "
  "lower end of what is desirable for this kind of study; we treat the three-seed "
  "mean as indicative, and a larger set would tighten the estimates.")

# ---------------------------------------------------------------- 6 CONCLUSION
H("6. Conclusion", 1)
P("We studied dynamic switch-to-controller assignment for distributed SDN in "
  "CitiVerse networks and showed, by direct measurement, that a learned policy can "
  "substantially improve control-plane responsiveness when the load is unevenly "
  "distributed across the city. The policy, a Dueling Double Deep Q-Network trained "
  "with prioritized replay, n-step returns and a curriculum, was learned in a fast "
  "analytical environment and transferred to a live Mininet and Ryu emulation. On the "
  "emulation, flow-setup latency under the skewed morning profile fell from 419 ms "
  "with zone-static assignment to 38 ms with the learned policy, and the "
  "per-controller measurements showed that the improvement comes from spreading "
  "control load away from the saturated residential controllers. The policy stayed "
  "competitive in the more balanced profiles.")
P("Beyond the numbers, the study keeps a fast analytical training environment "
  "separate from an empirical, packet-level evaluation, so the reported results rest "
  "on what the network does rather than on assumed parameter trends. This separation "
  "also keeps the main cost — the need to pre-train the agent — manageable, since the "
  "training environment can be extended cheaply. The natural next steps follow "
  "directly from the limitations above: the business-profile gap calls for richer "
  "coverage of moderate-load profiles during training, the single saturated-controller "
  "setting calls for measured controller-failure scenarios, and the spread across "
  "seeds calls for more repetitions to tighten the estimates.")

# ---------------------------------------------------------------- REFERENCES
H("References", 1)
refs = [
 # [1]
 "S. Islam, Z. A. Abdulsalam, B. A. Kumar, M. K. Hasan, R. Kolandaisamy and "
 "N. Safie, “Mobile networks toward 5G/6G: network architecture, opportunities and "
 "challenges in smart city,” IEEE Open J. Commun. Soc., vol. 6, pp. 3082–3093, 2025.",
 # [2]
 "T. Singh, A. Solanki, S. K. Sharma, A. Nayyar and A. Paul, “A decade review on "
 "smart cities: paradigms, challenges and opportunities,” IEEE Access, vol. 10, "
 "pp. 68319–68364, 2022.",
 # [3]
 "A. Volkov, A. Khakimov, A. Muthanna, R. Kirichek, A. Vladyko and "
 "A. Koucheryavy, “Interaction of the IoT traffic generated by a smart city "
 "segment with SDN core network,” in Wired/Wireless Internet Communications "
 "(WWIC 2017), Lecture Notes in Computer Science, vol. 10372, Springer, 2017, "
 "pp. 115–126.",
 # [4]
 "D. Kreutz et al., “Software-defined networking: a comprehensive survey,” "
 "Proc. IEEE, vol. 103, no. 1, pp. 14–76, 2015.",
 # [5]
 "M. Rahouti, K. Xiong and Y. Xin, “Secure software-defined networking "
 "communication systems for smart cities: current status, challenges, and "
 "trends,” IEEE Access, vol. 9, pp. 12083–12113, 2021.",
 # [6]
 "W. M. Othman et al., “Key enabling technologies for 6G,” J. Sens. Actuator "
 "Netw., vol. 14, no. 2, art. 30, 2025.",
 # [7]
 "Y. He, F. R. Yu, N. Zhao, V. C. M. Leung and H. Yin, “Software-defined networks "
 "with mobile edge computing and caching for smart cities: a big data deep "
 "reinforcement learning approach,” IEEE Commun. Mag., vol. 55, no. 12, "
 "pp. 31–37, 2017.",
 # [8]
 "G. Wang, Y. Zhao, J. Huang, Q. Duan and J. Li, “A k-means-based network "
 "partition algorithm for controller placement in SDN,” in Proc. IEEE ICC, 2016, "
 "pp. 1–6.",
 # [9]
 "H. Mostafaei, M. Menth and M. S. Obaidat, “A learning-automaton-based "
 "controller placement algorithm for software-defined networks,” in Proc. IEEE "
 "GLOBECOM, 2018, pp. 1–6.",
 # [10]
 "A. A. Ateya, A. Muthanna, A. Vybornova, A. D. Algarni, A. Abuarqoub et al., "
 "“Chaotic salp swarm algorithm for SDN multi-controller networks,” Engineering "
 "Science and Technology, an International Journal, vol. 22, no. 4, "
 "pp. 1001–1012, 2019.",
 # [11]
 "P. T. A. Quang, Y. Hadjadj-Aoul and A. Outtagarts, “A deep reinforcement "
 "learning approach for VNF forwarding graph embedding,” IEEE Trans. Netw. Serv. "
 "Manag., vol. 16, no. 4, pp. 1318–1331, 2019.",
 # [12]
 "V. Mnih et al., “Human-level control through deep reinforcement learning,” "
 "Nature, vol. 518, pp. 529–533, 2015.",
 # [13]
 "Z. Wang et al., “Dueling network architectures for deep reinforcement "
 "learning,” in Proc. ICML, 2016, pp. 1995–2003.",
 # [14]
 "H. van Hasselt, A. Guez and D. Silver, “Deep reinforcement learning with double "
 "Q-learning,” in Proc. AAAI, 2016, pp. 2094–2100.",
 # [15]
 "T. Schaul, J. Quan, I. Antonoglou and D. Silver, “Prioritized experience "
 "replay,” in Proc. ICLR, 2016.",
 # [16]
 "N. McKeown et al., “OpenFlow: enabling innovation in campus networks,” ACM "
 "SIGCOMM Comput. Commun. Rev., vol. 38, no. 2, pp. 69–74, 2008.",
 # [17]
 "B. Lantz, B. Heller and N. McKeown, “A network in a laptop: rapid prototyping "
 "for software-defined networks,” in Proc. ACM SIGCOMM HotNets, 2010.",
]
for i, rr in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.add_run(f"[{i}] ").bold = True
    p.add_run(rr)
    for run in p.runs: run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("SAVED:", OUT)
print("paragraphs:", len(doc.paragraphs))
